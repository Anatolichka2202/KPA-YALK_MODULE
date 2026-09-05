from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "output" / "documents" / "Непокрытые_проверки_УБСИ.docx"


def shade(cell, color):
    props = cell._tc.get_or_add_tcPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), color)
    props.append(element)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Arial"
    r._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.7)
section.bottom_margin = Cm(1.6)
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.6)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
styles["Normal"].font.size = Pt(10)
styles["Normal"].paragraph_format.space_after = Pt(5)
for name, size, color in [("Title", 21, "17365D"), ("Heading 1", 15, "17365D"), ("Heading 2", 12, "2F5597")]:
    styles[name].font.name = "Arial"
    styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles[name].font.size = Pt(size)
    styles[name].font.color.rgb = RGBColor.from_string(color)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("НЕПОКРЫТЫЕ ПРОВЕРКИ УБСИ\nИ ТРЕБОВАНИЯ К ДООСНАЩЕНИЮ СТЕНДА")
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Техническая записка к минимальной поставке ПО проверки ЯЛК-96 + ЯТП")
run.bold = True
run.font.size = Pt(12)

meta = doc.add_table(rows=4, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta.style = "Table Grid"
for row, (key, value) in enumerate([
    ("Объект", "УБСИ ЛВРМ.468157.002"),
    ("Основание", "Проект ТУ УБСИ; текущая конфигурация стенда КТМА"),
    ("Дата", "05.09.2026"),
    ("Статус", "Для согласования с производством и разработчиком оснастки"),
]):
    set_cell_text(meta.cell(row, 0), key, True, "FFFFFF")
    shade(meta.cell(row, 0), "2F5597")
    set_cell_text(meta.cell(row, 1), value)

doc.add_heading("1. Назначение", level=1)
doc.add_paragraph(
    "Документ фиксирует требования ТУ, которые текущая минимальная поставка стендового ПО "
    "не подтверждает автоматически. Отсутствие проверки нельзя заменять результатом «НОРМА». "
    "До появления подтверждённой методики и оснастки в отчёте должен применяться статус «НЕ ПРОВЕРЕНО»."
)

doc.add_heading("2. Уже закрытый автоматизированный объём", level=1)
for item in [
    "проверка ЯЛК-96: калибровка, 80 адресов, аналоговые значения, контактные состояния, обрыв, перегрузка ±12 В и безопасный сброс;",
    "проверка ЯТП: 30 каналов при ручных точках магазина 0 / 120 / 240 Ом;",
    "питание УБСИ: готовность не более 30 с, работа при 24 / 27 / 35 В, ток до 400 мА, выдержки 19 В — 5 мин и 37 В — 1 мин;",
    "контроль эталонного воздействия 6,2 ±0,03 В прибором В7-78/1;",
    "опциональная отдельная проверка контактных порогов ЯЛК 1,0 / 2,4 В.",
]:
    add_bullet(doc, item)

doc.add_page_break()
doc.add_heading("3. Требования, не проверяемые текущим стендом", level=1)

rows = [
    ("Питание датчиков 6,2 ±0,2 В при 350 мА",
     "Нет подтверждённого нагрузочного маршрута, распиновки и безопасной коммутации нагрузки.",
     "Электронная нагрузка либо эквивалент 17,7 Ом не менее 3 Вт; схема подключения; номера разъёмов/контактов; порядок включения.",
     "НЕ ПРОВЕРЕНО"),
    ("Защита при 450 мА и восстановление",
     "Не определены нагрузочный маршрут, критерий срабатывания и момент/условия восстановления.",
     "Электронная нагрузка либо эквивалент 13,8 Ом не менее 4 Вт; подтверждённый предел; выдержка; критерий отключения и восстановления.",
     "НЕ ПРОВЕРЕНО"),
    ("Входной ток канала ≤2 мкА",
     "АКИП измеряет суммарный ток всего УБСИ. Его показание не выделяет ток отдельного входного канала.",
     "Рассечка конкретного входа; микроамперметр/электрометр с подходящим диапазоном; схема, напряжение проверки и перечень каналов.",
     "НЕ ПРОВЕРЕНО"),
    ("Сопротивление изоляции ≥20 МОм",
     "У текущего стенда нет подтверждённого источника испытательного напряжения и измерителя с необходимым диапазоном и запасом.",
     "Мегаомметр/измеритель изоляции на установленное методикой напряжение (предварительно принято 100 В); карта точек; разряд и меры безопасности.",
     "НЕ ПРОВЕРЕНО"),
]
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Требование", "Почему сейчас нельзя подтвердить", "Что нужно предоставить", "Статус"]
for i, text in enumerate(headers):
    set_cell_text(table.cell(0, i), text, True, "FFFFFF", 8)
    shade(table.cell(0, i), "2F5597")
for requirement, reason, needed, status in rows:
    cells = table.add_row().cells
    for index, value in enumerate((requirement, reason, needed, status)):
        set_cell_text(cells[index], value, bold=index == 0, size=8)
    set_cell_text(cells[3], status, True, "C00000", 8)
for row in table.rows:
    row.cells[0].width = Cm(3.5)
    row.cells[1].width = Cm(5.0)
    row.cells[2].width = Cm(6.6)
    row.cells[3].width = Cm(2.4)

doc.add_heading("4. Что требуется от производства", level=1)
for item in [
    "подтверждённая схема подключения для каждого требования: разъём, номера контактов, полярность, общий провод и допустимый ток;",
    "утверждённая последовательность действий и точные критерии «НОРМА / НЕ НОРМА»;",
    "перечень средств измерений, диапазоны, класс точности и требования к поверке;",
    "решение, какие каналы проверяются полностью, а какие принимаются по производственному контролю ячейки;",
    "разрешение на автоматическое управление нагрузкой либо форма ручного протокола.",
]:
    add_bullet(doc, item)

doc.add_page_break()
doc.add_heading("5. Допустимый временный ручной ввод", level=1)
doc.add_paragraph(
    "До автоматизации допускается заносить результат внешнего измерения в будущий регистратор, "
    "но только как ссылку на самостоятельный протокол. Простого флажка оператора недостаточно."
)
manual = doc.add_table(rows=1, cols=2)
manual.style = "Table Grid"
for i, text in enumerate(("Обязательное поле", "Содержание")):
    set_cell_text(manual.cell(0, i), text, True, "FFFFFF")
    shade(manual.cell(0, i), "2F5597")
for key, value in [
    ("Требование ТУ", "номер пункта и наименование параметра"),
    ("Измеренное значение", "значение, единица, нижний/верхний предел"),
    ("Схема проверки", "разъём, контакты, подключённая нагрузка или рассечка"),
    ("Средство измерений", "тип, заводской номер, дата/срок поверки"),
    ("Прослеживаемость", "номер внешнего протокола, дата, оператор"),
    ("Решение", "НОРМА / НЕ НОРМА; «НЕ ПРОВЕРЕНО» при неполных данных"),
]:
    cells = manual.add_row().cells
    set_cell_text(cells[0], key, True)
    set_cell_text(cells[1], value)

doc.add_heading("6. Правило интеграции в ПО", level=1)
for item in [
    "Автоматическая процедура добавляется в обязательный сценарий только после подтверждения схемы и безопасного прогона на живом стенде.",
    "Ручной результат считается доказательством только при заполнении всех полей раздела 5 и хранении внешнего протокола.",
    "Если оборудование или инструкция не предоставлены, краткий отчёт ТУ явно перечисляет непроверенные пункты и не выдаёт заключение о полном соответствии УБСИ.",
    "Климатические испытания, масса, габариты, маркировка и внешний вид ведутся отдельным регистратором; после климатического воздействия стенд повторяет обычную электрическую проверку.",
]:
    add_bullet(doc, item)

doc.add_heading("7. Решение на текущую поставку", level=1)
p = doc.add_paragraph()
r = p.add_run("В обязательную поставку входят ЯЛК-96, ЯТП и реализованный цикл питания. ")
r.bold = True
p.add_run(
    "Питание датчиков 350/450 мА, входной ток ≤2 мкА и изоляция ≥20 МОм "
    "остаются «НЕ ПРОВЕРЕНО» до получения оборудования, оснастки и утверждённой инструкции."
)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("УБСИ ЛВРМ.468157.002 · техническая записка · 05.09.2026")
fr.font.name = "Arial"
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor(100, 110, 120)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
